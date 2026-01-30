import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from document.pdf import generate_smart_filename

HAS_DOCX = True

# --- Converter Logic ---

FIELD_RE = re.compile(r"\{\{FIELD:([^|}]+)(?:\|([^}]+))?\}\}")
TABLE_START_RE = re.compile(r"\{\{TABLE:([^}]+)\}\}")
TABLE_END = "{{ENDTABLE}}"
CELL_RE = re.compile(r"\{\{CELL:([^:}]+):(\d+):(\d+)\}\}")

def parse_attrs(attr_string):
    """Parse attribute string like 'label=Nom|lines=1|value=foo'"""
    res = {}
    if not attr_string:
        return res
    parts = attr_string.split("|")
    for p in parts:
        if "=" in p:
            k, v = p.split("=", 1)
            res[k.strip()] = v.strip()
    return res

def add_underline_paragraph(doc, lines=1, value=None):
    """Add a paragraph with underline or value text"""
    if value:
        p = doc.add_paragraph(value)
        return p
    
    # Create underline using bottom border
    p = doc.add_paragraph()
    p.add_run("_" * 50)  # Fallback: underscore line
    return p

def add_heading(doc, line):
    """Add heading based on markdown level"""
    if line.startswith("# "):
        level = 0
        text = line.lstrip("# ").strip()
    elif line.startswith("## "):
        level = 1
        text = line.lstrip("## ").strip()
    else:
        level = 2
        text = line.lstrip("### ").strip()
    
    h = doc.add_heading(text, level=level)
    return h

def add_bullet(doc, text):
    """Add bullet point"""
    p = doc.add_paragraph(text.strip(), style='List Bullet')
    return p

def process_markdown_formatting(paragraph, text):
    """Process inline markdown (bold, italic) and add to paragraph"""
    # Split by bold and italic markers
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            # Italic
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            # Regular text
            paragraph.add_run(part)

def convert_markdown_to_docx(md_text, output_path):
    """
    Convert evaluation markdown with field markers to DOCX.
    
    Args:
        md_text: Markdown text with {{FIELD:...}} and {{TABLE:...}} markers
        output_path: Path to save the DOCX file
    """
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    lines = md_text.splitlines()
    i = 0
    n = len(lines)
    
    while i < n:
        line = lines[i].rstrip()
        
        # Skip YAML frontmatter
        if i == 0 and line.strip() == '---':
            i += 1
            while i < n and lines[i].strip() != '---':
                i += 1
            i += 1
            continue
        
        # Headings
        if line.startswith("#"):
            add_heading(doc, line)
            i += 1
            continue
        
        # Table start
        if line.strip().startswith("{{TABLE:"):
            m = TABLE_START_RE.match(line.strip())
            if not m:
                i += 1
                continue
            
            table_info = m.group(1)
            attrs = parse_attrs(table_info)
            cols = int(attrs.get("cols", "2"))
            rows = int(attrs.get("rows", "1"))
            
            # Collect cell markers and content
            cell_text = {}
            j = i + 1
            while j < n and lines[j].strip() != TABLE_END:
                # Find cell markers
                for cm in CELL_RE.finditer(lines[j]):
                    tabid, r, c = cm.group(1), int(cm.group(2)), int(cm.group(3))
                    after = lines[j][cm.end():].strip()
                    cell_text[(r, c)] = after
                j += 1
            
            # Create DOCX table
            table = doc.add_table(rows=rows + 1, cols=cols)
            table.style = 'Light Grid Accent 1'
            
            # Find header row from markdown table syntax
            k = i + 1
            header_line = None
            while k < j:
                if '|' in lines[k] and not lines[k].strip().startswith('{{'):
                    header_line = lines[k]
                    break
                k += 1
            
            if header_line:
                # Parse header
                parts = [p.strip() for p in header_line.strip().strip("|").split("|")]
                for ci, text in enumerate(parts):
                    if ci < cols:
                        table.cell(0, ci).text = text
            
            # Fill cell placeholders
            for r in range(1, rows + 1):
                for c in range(1, cols + 1):
                    text = cell_text.get((r, c), "")
                    if r <= rows and c <= cols:
                        table.cell(r, c - 1).text = text
            
            i = j + 1
            continue
        
        # Field markers inline
        fm = FIELD_RE.search(line)
        if fm:
            before = line[:fm.start()].strip()
            fid = fm.group(1)
            attrs = parse_attrs(fm.group(2))
            lines_count = int(attrs.get("lines", "1"))
            value = attrs.get("value", None)
            
            p = doc.add_paragraph()
            if before:
                process_markdown_formatting(p, before + " ")
            
            # Insert underline or value
            if value:
                p.add_run(value)
            else:
                p.add_run("_" * 50)
            
            # Add remaining text
            after = line[fm.end():].strip()
            if after:
                p.add_run(" " + after)
            
            i += 1
            continue
        
        # Bullet points
        if line.strip().startswith("- "):
            add_bullet(doc, line.strip()[2:])
            i += 1
            continue
        
        # Horizontal rule
        if line.strip() == "---":
            doc.add_paragraph()  # Just add spacing
            i += 1
            continue
        
        # Empty line
        if line.strip() == "":
            doc.add_paragraph()
            i += 1
            continue
        
        # Regular paragraph
        p = doc.add_paragraph()
        process_markdown_formatting(p, line)
        i += 1
    
    # Save document
    doc.save(output_path)
    return output_path

# --- Export Functions ---

def save_fiche_to_docx(content, lesson_topic, class_level, output_dir, queue):
    if not HAS_DOCX:
        queue.put(("log", "❌ DOCX export requires python-docx. Run: pip install python-docx"))
        return None

    os.makedirs(output_dir, exist_ok=True)
    
    # Create smart filename that doesn't overwrite
    filename = generate_smart_filename("Fiche", lesson_topic, class_level, output_dir, "docx")
    full_path = os.path.join(output_dir, filename)

    try:
        doc = Document()
        # Set font properties on the document's default style
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        lines = content.splitlines()
        i = 0
        while i < len(lines):
            line = lines[i].rstrip()
            i += 1
            
            if not line:
                continue
                
            # Main title
            if line.startswith('# '):
                doc.add_heading(line[2:].strip(), level=0)
                continue
                
            # Section heading
            if line.startswith('## '):
                doc.add_heading(line[3:].strip(), level=1)
                continue
                
            # Subheading or metadata
            if line.startswith('### '):
                heading_text = line[4:].strip()
                # Check if it's metadata format: ### **Key**: Value
                m_meta = re.match(r'\*\*([^*]+)\*\*\s*:\s*(.*)$', heading_text)
                if m_meta:
                    p = doc.add_paragraph()
                    run_b = p.add_run(m_meta.group(1) + ": ")
                    run_b.bold = True
                    run_b.font.color.rgb = RGBColor(0, 102, 204)
                    p.add_run(m_meta.group(2))
                else:
                    doc.add_heading(heading_text, level=2)
                continue
                
            # Bullet points
            if line.startswith('- ') or line.startswith('* '):
                bullet_text = line[2:].strip()
                # Handle inline markdown
                bullet_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', bullet_text)  # Bold (can't do in list easily)
                doc.add_paragraph(bullet_text, style='List Bullet')
                continue
            
            # Numbered lists
            if re.match(r'^\d+\.\s', line):
                number_text = re.sub(r'^\d+\.\s', '', line)
                number_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', number_text)
                doc.add_paragraph(number_text, style='List Number')
                continue
            
            # Tables (check if line contains |)
            if '|' in line and line.strip().startswith('|'):
                # Collect table lines
                table_lines = [line]
                while i < len(lines) and '|' in lines[i]:
                    table_lines.append(lines[i].rstrip())
                    i += 1
                
                # Parse table
                table_data = []
                for tline in table_lines:
                    # Skip separator rows (| --- | --- |)
                    if re.match(r'\|\s*[-:]+\s*\|', tline):
                        continue
                    cells = [cell.strip() for cell in tline.strip('|').split('|')]
                    if cells:
                        table_data.append(cells)
                
                if table_data:
                    # Create table
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    table.style = 'Light Grid Accent 1'
                    
                    for row_idx, row_data in enumerate(table_data):
                        for col_idx, cell_text in enumerate(row_data):
                            if col_idx < len(table.rows[row_idx].cells):
                                table.rows[row_idx].cells[col_idx].text = cell_text
                                # Bold first row (header)
                                if row_idx == 0:
                                    for paragraph in table.rows[row_idx].cells[col_idx].paragraphs:
                                        for run in paragraph.runs:
                                            run.bold = True
                continue
            
            # Regular paragraph with inline markdown
            if line.strip():
                p = doc.add_paragraph()
                # Process inline bold and italic
                parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    elif part.startswith('*') and part.endswith('*'):
                        run = p.add_run(part[1:-1])
                        run.italic = True
                    else:
                        p.add_run(part)

        doc.save(full_path)
        queue.put(("log", f"💾 DOCX saved: {full_path}"))
        return full_path
    except Exception as e:
        queue.put(("log", f"❌ DOCX Save Error: {e}"))
        import traceback
        queue.put(("log", f"Traceback: {traceback.format_exc()}"))
        return None

def save_evaluation_to_docx(content, topics_list, class_level, output_dir, queue):
    """
    Save evaluation content as DOCX using the special converter.
    This preserves {{FIELD:...}}, {{TABLE:...}}, and {{CELL:...}} markers
    for editable evaluations in Word.
    """
    if not HAS_DOCX:
        queue.put(("log", "❌ DOCX export requires python-docx. Run: pip install python-docx"))
        return None

    os.makedirs(output_dir, exist_ok=True)
    
    # Create smart filename for evaluation
    topics_text = "_".join(topics_list[:2])  # Use first 2 topics
    if len(topics_list) > 2:
        topics_text += "_etc"
    filename = generate_smart_filename("Eval", topics_text, class_level, output_dir, "docx")
    full_path = os.path.join(output_dir, filename)

    try:
        # Use the special converter that handles field markers
        convert_markdown_to_docx(content, full_path)
        queue.put(("log", f"💾 Evaluation DOCX saved: {full_path}"))
        queue.put(("log", "ℹ️  Field markers ({{FIELD:...}}, {{TABLE:...}}) preserved for editing in Word"))
        return full_path
    except Exception as e:
        queue.put(("log", f"❌ Evaluation DOCX Save Error: {e}"))
        import traceback
        queue.put(("log", f"Traceback: {traceback.format_exc()}"))
        return None
